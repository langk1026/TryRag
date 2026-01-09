import io
from PyPDF2 import PdfReader
from docx import Document
import openpyxl
from backend.core.logger import setup_logger

logger = setup_logger(__name__)


class DocumentProcessor:
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.txt', '.xlsx', '.md']

    def extract_text(self, content, file_name):
        extension = self._get_extension(file_name)

        if extension not in self.supported_extensions:
            logger.warning(f"Unsupported file type: {extension} for {file_name}")
            return ""

        try:
            if extension == '.pdf':
                return self._extract_from_pdf(content)
            elif extension == '.docx':
                return self._extract_from_docx(content)
            elif extension == '.txt' or extension == '.md':
                return self._extract_from_text(content)
            elif extension == '.xlsx':
                return self._extract_from_excel(content)
            else:
                return ""
        except Exception as e:
            logger.error(f"Failed to extract text from {file_name}: {str(e)}")
            return ""

    def _get_extension(self, file_name):
        return '.' + file_name.split('.')[-1].lower()

    def _extract_from_pdf(self, content):
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            text_parts.append(page_text)

        full_text = '\n\n'.join(text_parts)
        logger.debug(f"Extracted {len(full_text)} characters from PDF")
        return full_text

    def _extract_from_docx(self, content):
        doc_file = io.BytesIO(content)
        document = Document(doc_file)

        text_parts = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        full_text = '\n\n'.join(text_parts)
        logger.debug(f"Extracted {len(full_text)} characters from DOCX")
        return full_text

    def _extract_from_text(self, content):
        text = content.decode('utf-8', errors='ignore')
        logger.debug(f"Extracted {len(text)} characters from text file")
        return text

    def _extract_from_excel(self, content):
        excel_file = io.BytesIO(content)
        workbook = openpyxl.load_workbook(excel_file, data_only=True)

        text_parts = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"Sheet: {sheet_name}")

            for row in sheet.iter_rows(values_only=True):
                row_text = '\t'.join([str(cell) if cell is not None else '' for cell in row])
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = '\n'.join(text_parts)
        logger.debug(f"Extracted {len(full_text)} characters from Excel")
        return full_text

    def extract_text_with_pages(self, content, file_name):
        extension = self._get_extension(file_name)

        if extension not in self.supported_extensions:
            logger.warning(f"Unsupported file type: {extension} for {file_name}")
            return []

        try:
            if extension == '.pdf':
                return self._extract_from_pdf_with_pages(content)
            elif extension == '.docx':
                return self._extract_from_docx_with_pages(content)
            elif extension == '.txt' or extension == '.md':
                return self._extract_from_text_with_pages(content)
            elif extension == '.xlsx':
                return self._extract_from_excel_with_pages(content)
            else:
                return []
        except Exception as e:
            logger.error(f"Failed to extract text with pages from {file_name}: {str(e)}")
            return []

    def _extract_from_pdf_with_pages(self, content):
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)

        pages_data = []
        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text()
            if page_text.strip():
                pages_data.append({
                    'page_number': page_num,
                    'text': page_text
                })

        logger.debug(f"Extracted {len(pages_data)} pages from PDF")
        return pages_data

    def _extract_from_docx_with_pages(self, content):
        doc_file = io.BytesIO(content)
        document = Document(doc_file)

        page_text = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                page_text.append(paragraph.text)

        pages_data = [{
            'page_number': 1,
            'text': '\n\n'.join(page_text)
        }]

        logger.debug(f"Extracted 1 page from DOCX")
        return pages_data

    def _extract_from_text_with_pages(self, content):
        text = content.decode('utf-8', errors='ignore')

        pages_data = [{
            'page_number': 1,
            'text': text
        }]

        logger.debug(f"Extracted 1 page from text file")
        return pages_data

    def _extract_from_excel_with_pages(self, content):
        excel_file = io.BytesIO(content)
        workbook = openpyxl.load_workbook(excel_file, data_only=True)

        pages_data = []
        for page_num, sheet_name in enumerate(workbook.sheetnames, start=1):
            sheet = workbook[sheet_name]
            text_parts = [f"Sheet: {sheet_name}"]

            for row in sheet.iter_rows(values_only=True):
                row_text = '\t'.join([str(cell) if cell is not None else '' for cell in row])
                if row_text.strip():
                    text_parts.append(row_text)

            pages_data.append({
                'page_number': page_num,
                'text': '\n'.join(text_parts)
            })

        logger.debug(f"Extracted {len(pages_data)} sheets from Excel")
        return pages_data
